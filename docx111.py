from docx import Document

path = "E:\\Desktop\\PyCode\\temp01.docx"

document = Document()
document.add_heading('玉米行情周报', 0)  #插入标题
p = document.add_paragraph('A plain paragraph having some ')   #插入段落
p.add_run('bold').bold = True
p.add_run(' and some ')
p.add_run('italic.').italic = True

document.add_heading('Heading, level 1', level=1)
document.add_paragraph('Intense quote', style='IntenseQuote')

document.add_paragraph('first item in unordered list', style='ListBullet')
document.add_paragraph('first item in ordered list', style='ListNumber')
document.add_picture('http://image.nmc.cn/product/2019/07/04/AMDF/medium/SEVP_NMC_AMDF_SFER_EDRF_ACHN_L88_P9_20190704000000000.jpg?v=1562204332470') #插入图片


document.save(path)